"""
DocxBuilder — Word (.docx) Report Renderer
============================================
Concrete implementation of ReportBuilder using python-docx.

Design decisions:
  - Uses a custom UAV-CD-APP document style (blue accent, Calibri body).
  - Equations use Word's native OMML math objects (m:oMath) inside a single
    paragraph with a right-aligned tab stop for the equation number.
    This produces real Word equation objects — editable via Word's built-in
    equation editor — without any table wrapper that might show borders.
  - Equation numbers are auto-generated as (section.subnum) — sections call
    ``add_equation(latex_src)`` with no number argument; the renderer
    assigns and resets the sub-counter on each level-1 heading.
  - Inline math: anywhere text is rendered (add_paragraph, add_note,
    key-value list values, table cells), the sigil ``$${LaTeX}$$`` is
    detected and converted to an inline OMML run.  Example:
        rb.add_paragraph("The efficiency $${\\eta_p}$$ captures losses.")
  - Tables use 'Table Grid' style with blue header row and alternating rows.
  - Figures are written to a NamedTemporaryFile, inserted, then cleaned up.
  - All widths are specified in centimetres (docx uses EMU internally).
"""

from __future__ import annotations

import io
import math
import re
import tempfile
import os
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from lxml import etree

from app.core.reports.renderer import ReportBuilder, ReportConfig
from app.core.reports.renderers.latex_omml import latex_to_omml

# Brand colours
_BLUE   = RGBColor(0x1A, 0x5C, 0x96)   # heading blue
_DARK   = RGBColor(0x2C, 0x2C, 0x2C)   # body text (near-black)
_GREY   = RGBColor(0xF0, 0xF0, 0xF0)   # note background
_ACCENT = RGBColor(0x27, 0xAE, 0x60)   # green accent for values
_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# OMML (Office Math Markup Language) namespace
_MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_M = f"{{{_MATH_NS}}}"


def _shade_cell(cell, fill_hex: str) -> None:
    """Apply a solid background shade to a table cell via OOXML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _make_omath(equation_src: str) -> etree._Element:
    """
    Build a structured OMML ``<m:oMath>`` element from a LaTeX math source.

    The input is LaTeX (e.g. ``r"\\frac{W_E}{W_{TO}} = a \\cdot W_{TO}^{b}"``).
    The output has real OMML structure — ``<m:f>`` for fractions, ``<m:rad>``
    for square roots, ``<m:sSup>``/``<m:sSub>``/``<m:sSubSup>`` for sub/sup,
    ``<m:d>`` for delimiters — so Word renders it as a real equation object
    (click → equation editor opens with proper math layout).
    """
    return latex_to_omml(equation_src)


# Inline math sigil. Any occurrence of ``$${LATEX}$$`` inside a text
# argument (paragraph, note, table cell, kv-value, …) is replaced by a real
# inline OMML run. Triple-brace ensures the sigil won't collide with normal
# prose. The pattern is non-greedy and DOTALL so newlines inside the math
# source are allowed.
_INLINE_MATH = re.compile(r"\$\$\{(.+?)\}\$\$", re.DOTALL)


def _emit_text_with_math(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    complex_script_font: Optional[str] = None,
) -> None:
    """Append runs to ``paragraph`` for ``text``, converting any
    ``$${LaTeX}$$`` segments to inline OMML math.

    Plain text segments become regular ``<w:r>`` runs (honouring
    bold/italic); math segments become inline ``<m:oMath>`` elements
    appended directly to the paragraph's ``<w:p>``.

    If ``complex_script_font`` is set, it is applied to each text run via
    ``run.font.complex_script_font`` — needed for Arabic glyphs since the
    Latin body font (Calibri) lacks proper Arabic shaping.
    """
    pos = 0
    for m in _INLINE_MATH.finditer(text):
        before = text[pos:m.start()]
        if before:
            run = paragraph.add_run(before)
            run.bold = bold
            run.italic = italic
            if complex_script_font:
                _apply_complex_script_font(run, complex_script_font)
        # Append the OMML element directly to the <w:p> (not wrapped in a run)
        paragraph._p.append(latex_to_omml(m.group(1)))
        pos = m.end()
    tail = text[pos:]
    if tail:
        run = paragraph.add_run(tail)
        run.bold = bold
        run.italic = italic
        if complex_script_font:
            _apply_complex_script_font(run, complex_script_font)


def _apply_complex_script_font(run, font_name: str) -> None:
    """Set the complex-script font on a run via ``<w:rFonts w:cs="…"/>``."""
    rPr = run._r.get_or_add_rPr()
    # python-docx exposes rFonts via descriptor but only for the Latin face;
    # set the cs (complex-script) attribute directly.
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), font_name)


def _set_paragraph_rtl(paragraph) -> None:
    """Mark a paragraph as right-to-left via ``<w:bidi/>``.

    Affects default alignment, tab-stop direction, and list ordering for
    Arabic body text.
    """
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))


def _set_table_rtl(table) -> None:
    """Mark a table as RTL via ``<w:bidiVisual/>`` on the table properties.

    Word will display columns right-to-left (the first authored column
    becomes the rightmost visible column).
    """
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))


class DocxBuilder(ReportBuilder):
    """Word (.docx) implementation of the ReportBuilder interface."""

    def __init__(self, config: ReportConfig) -> None:
        self._config = config
        self._doc = Document()
        self._section_counter: int = 0
        self._figure_counter:  int = 0
        self._table_counter:   int = 0
        self._eq_subcounter:   int = 0   # resets on every level-1 heading
        # Language-driven layout switches. None == LTR / no special font.
        self._is_rtl: bool = config.language.is_rtl
        self._cs_font: Optional[str] = (
            config.language.complex_script_font if self._is_rtl else None
        )
        self._setup_document()

    # ── Document setup ────────────────────────────────────────────────────

    def _setup_document(self) -> None:
        """Configure page margins and override default paragraph styles."""
        # Page margins
        for section in self._doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.5)

        # Default body font
        normal = self._doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = _DARK

        # Heading 1
        h1 = self._doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = _BLUE

        # Heading 2
        h2 = self._doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = _BLUE

        # Heading 3
        h3 = self._doc.styles["Heading 3"]
        h3.font.name = "Calibri"
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.color.rgb = _DARK

    # ── ReportBuilder implementation ──────────────────────────────────────

    def add_heading(self, text: str, level: int = 1) -> None:
        if level == 1:
            self._section_counter += 1
            self._eq_subcounter = 0   # restart equation numbering per section
            # Section/heading numbers stay Western (Latin digits) even in RTL
            # mode — they're cross-reference IDs, not localised content.
            numbered = f"{self._section_counter}.  {text}"
        else:
            numbered = text
        heading = self._doc.add_heading(numbered, level=min(level, 3))
        # Headings carry text-with-math too via the inline-math sigil, but
        # python-docx already inserted the literal text as a run. Apply RTL +
        # complex-script font to the heading paragraph and its existing runs.
        if self._is_rtl:
            _set_paragraph_rtl(heading)
            if self._cs_font:
                for run in heading.runs:
                    _apply_complex_script_font(run, self._cs_font)

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        indent: bool = False,
    ) -> None:
        p = self._doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(1)
        if self._is_rtl:
            _set_paragraph_rtl(p)
        _emit_text_with_math(
            p, text, bold=bold, italic=italic,
            complex_script_font=self._cs_font,
        )

    def add_page_break(self) -> None:
        self._doc.add_page_break()

    def add_table(
        self,
        headers: list[str],
        rows: list[list[str]],
        caption: str = "",
        style: str = "default",
    ) -> None:
        self._table_counter += 1
        col_count = len(headers)
        row_count = 1 + len(rows)

        table = self._doc.add_table(rows=row_count, cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if self._is_rtl:
            _set_table_rtl(table)

        # Header row
        hdr_cells = table.rows[0].cells
        for i, hdr in enumerate(headers):
            cell = hdr_cells[i]
            _shade_cell(cell, "1A5C96")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if self._is_rtl:
                _set_paragraph_rtl(p)
            _emit_text_with_math(
                p, str(hdr), bold=True, complex_script_font=self._cs_font,
            )
            # White text on the blue header — recolor whatever runs ended up there
            for run in p.runs:
                run.font.color.rgb = _WHITE

        # Data rows
        for r_idx, row_data in enumerate(rows):
            cells = table.rows[r_idx + 1].cells
            fill = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, cell_text in enumerate(row_data):
                cell = cells[c_idx]
                _shade_cell(cell, fill)
                p = cell.paragraphs[0]
                # In RTL mode "first column" is rendered rightmost, so the
                # author's first column should remain LEFT-aligned (which
                # Word will visually flip to RIGHT in the bidi table).
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                if self._is_rtl:
                    _set_paragraph_rtl(p)
                _emit_text_with_math(
                    p, str(cell_text), complex_script_font=self._cs_font,
                )

        if caption:
            cap_p = self._doc.add_paragraph()
            if self._is_rtl:
                _set_paragraph_rtl(cap_p)
            # Table number stays Western (Latin digits) — it's a cross-ref ID.
            _emit_text_with_math(
                cap_p, f"Table {self._table_counter}: {caption}",
                italic=True, complex_script_font=self._cs_font,
            )
            for run in cap_p.runs:
                run.font.size = Pt(9)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._doc.add_paragraph()   # spacer

    def add_figure(
        self,
        image_bytes: bytes,
        caption: str = "",
        width_cm: float = 16.0,
    ) -> None:
        self._figure_counter += 1
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        try:
            tmp.write(image_bytes)
            tmp.flush()
            tmp.close()
            p = self._doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(tmp.name, width=Cm(width_cm))
        finally:
            os.unlink(tmp.name)

        if caption:
            cap_p = self._doc.add_paragraph()
            if self._is_rtl:
                _set_paragraph_rtl(cap_p)
            _emit_text_with_math(
                cap_p, f"Figure {self._figure_counter}: {caption}",
                italic=True, complex_script_font=self._cs_font,
            )
            for run in cap_p.runs:
                run.font.size = Pt(9)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._doc.add_paragraph()

    def add_equation(self, equation_text: str) -> None:
        """Add a displayed equation as a real Word OMML math object.

        Layout — single paragraph with:
            [indent][OMML math object][right-aligned tab][(section.sub)]

        The equation number is auto-generated as ``(section.sub)`` where
        ``section`` is the current level-1 heading count and ``sub`` is the
        equation sequence within that section (reset on every new section).

        The equation itself is a genuine Word equation object — click it to
        edit in Word's built-in equation editor.

        Sadraey references should appear once in the section's introductory
        paragraph (e.g. "follows Sadraey 2020, Sec. 2.9"), not under each
        equation.
        """
        self._eq_subcounter += 1
        eq_number = f"{self._section_counter}.{self._eq_subcounter}"

        # Build a paragraph that will hold the OMML + the right-aligned number
        p = self._doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()

        # Paragraph indent: ~1 cm left so the equation is visually inset
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "567")   # 1 cm ≈ 567 twips
        pPr.append(ind)

        # Right-aligned tab stop at ~15 cm for the equation number
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "8504")   # 15 cm ≈ 8504 twips
        tabs.append(tab)
        pPr.append(tabs)

        # Insert the OMML <m:oMath> element directly into the <w:p>
        p._p.append(_make_omath(equation_text))

        # Tab + number run
        tab_r = OxmlElement("w:r")
        tab_t = OxmlElement("w:tab")
        tab_r.append(tab_t)
        p._p.append(tab_r)

        num_r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")   # 10 pt
        rPr.append(sz)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "606060")
        rPr.append(color)
        num_r.append(rPr)
        num_t = OxmlElement("w:t")
        num_t.text = f"({eq_number})"
        num_r.append(num_t)
        p._p.append(num_r)

    def add_key_value_list(
        self,
        items: list[tuple[str, str]],
        columns: int = 2,
    ) -> None:
        """Render as a bordered two-column table (label | value).

        Both label and value support inline ``$${LaTeX}$$`` math.
        """
        if not items:
            return
        table = self._doc.add_table(rows=len(items), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        if self._is_rtl:
            _set_table_rtl(table)
        for r, (lbl, val) in enumerate(items):
            row = table.rows[r]
            fill = "F5F5F5" if r % 2 == 0 else "FFFFFF"
            _shade_cell(row.cells[0], fill)
            _shade_cell(row.cells[1], fill)
            for cell, content, bold in (
                (row.cells[0], str(lbl), True),
                (row.cells[1], str(val), False),
            ):
                p = cell.paragraphs[0]
                if self._is_rtl:
                    _set_paragraph_rtl(p)
                _emit_text_with_math(
                    p, content, bold=bold, complex_script_font=self._cs_font,
                )
        self._doc.add_paragraph()

    def add_bulleted_list(self, items: list[str]) -> None:
        for item in items:
            self._doc.add_paragraph(item, style="List Bullet")

    def add_note(self, text: str) -> None:
        p = self._doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1.0)
        p.paragraph_format.right_indent = Cm(1.0)
        if self._is_rtl:
            _set_paragraph_rtl(p)
        # "Note:" prefix label, then the message (with inline math).
        # The prefix word itself is not translated here — section authors
        # who want a localised "Note" prefix can include it in the text
        # argument via ctx.t("note.prefix") + their message.
        label = p.add_run("Note:  ")
        label.bold = True
        label.font.size = Pt(10)
        label.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        if self._cs_font:
            _apply_complex_script_font(label, self._cs_font)
        _emit_text_with_math(
            p, text, italic=True, complex_script_font=self._cs_font,
        )
        # Re-style the emitted runs so the note appears uniformly muted
        for run in p.runs[1:]:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    def add_horizontal_rule(self) -> None:
        p = self._doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1A5C96")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def save(self, output_path: str) -> None:
        self._doc.save(output_path)
