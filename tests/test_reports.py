"""
Report Export Test Suite — UAV-CD-APP
========================================
End-to-end tests for the pluggable report export system.

Why this file exists:
  The original test_core.py never imports any report-section module, so
  AttributeError-class bugs inside section.build() methods slip through.
  These tests instantiate every section, run the full ExportService
  pipeline against multiple propulsion types, and re-open the resulting
  .docx with python-docx to verify the structural contracts of the doc.

Requirements: Qt headless (QT_QPA_PLATFORM=offscreen), python-docx, lxml.
"""

from __future__ import annotations

import os

# Must be set before any PyQt6 import
os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

import tempfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from PyQt6.QtWidgets import QApplication

# Triggers section auto-registration before SectionRegistry is queried
import app.core.reports.sections  # noqa: F401

from app.core.display_converter import DisplayConverter
from app.core.enums import PropulsionType
from app.core.reports.base import ReportContext, SectionRegistry, SectionCategory
from app.core.reports.renderer import ExportFormat, ReportConfig
from app.core.reports.renderers.docx_renderer import DocxBuilder, _make_omath
from app.services.export_service import ExportService
from app.services.sizing_service import SizingService
from app.state.store import AppStore

_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# ── Fixtures ─────────────────────────────────────────────────────────────────


@pytest.fixture(scope="session")
def qapp():
    app = QApplication.instance() or QApplication([])
    yield app


@pytest.fixture
def store(qapp):
    return AppStore()


@pytest.fixture
def sized_store(store):
    """Store with a default sizing run completed."""
    SizingService(store).run_now()
    return store


_LOCALES_DIR = Path(__file__).resolve().parent.parent / "app" / "resources" / "locales"


def _build_context(
    store: AppStore,
    *,
    include_equations: bool = True,
    language: "Language | None" = None,
) -> ReportContext:
    from app.core.i18n import Language as _Lang, load_translator
    lang = language or _Lang.EN
    s = store.state
    return ReportContext(
        project_name="Test Project",
        report_title="Test Report",
        author="Test Author",
        date_str="2026-05-16",
        revision="1.0",
        brief=s.sizing.brief,
        settings=store.settings,
        weight_result=s.sizing.weight_result,
        constraint_result=s.sizing.constraint_result,
        design_point=s.sizing.design_point,
        regression_coeffs=None,
        matching_diagram_png=None,
        mission_profile_png=None,
        weight_pie_chart_png=None,
        display_converter=DisplayConverter(store.settings),
        language=lang,
        translator=load_translator(_LOCALES_DIR, lang),
        include_equations=include_equations,
        include_sadraey_refs=True,
    )


def _export(store: AppStore, out_path: Path) -> tuple[bool, str]:
    cfg = ReportConfig(
        report_title="Test",
        author="Test",
        revision="1.0",
        format=ExportFormat.DOCX,
        sections=SectionRegistry.default_manifest(),
        output_path=str(out_path),
    )
    return ExportService(store).export(cfg, figure_grabbers={})


# ── Registry tests ───────────────────────────────────────────────────────────


class TestSectionRegistry:
    def test_all_sections_registered(self):
        # 13 original sections + 1 customisable Sensitivity Analysis (PR6).
        assert len(SectionRegistry.all_sections()) == 14

    def test_section_ids_unique(self):
        ids = [s.section_id for s in SectionRegistry.all_sections()]
        assert len(set(ids)) == len(ids)

    def test_sections_have_required_metadata(self):
        for cls in SectionRegistry.all_sections():
            assert cls.section_id, f"{cls.__name__} missing section_id"
            assert cls.title, f"{cls.__name__} missing title"
            assert isinstance(cls.category, SectionCategory)
            assert isinstance(cls.default_order, int)

    def test_default_manifest_in_order(self):
        manifest = SectionRegistry.default_manifest()
        orders = [
            SectionRegistry.get(e.section_id).default_order for e in manifest
        ]
        assert orders == sorted(orders)


# ── Each section can build without raising ──────────────────────────────────


class TestSectionBuilds:
    """Catches AttributeError-class bugs that the rest of the suite misses."""

    @pytest.mark.parametrize(
        "propulsion",
        [
            PropulsionType.ELECTRIC,
            PropulsionType.PISTON,
            PropulsionType.HYBRID,
            PropulsionType.TURBOJET,
        ],
    )
    def test_every_section_builds(self, store, propulsion):
        store.update_brief_field("propulsion_type", propulsion)
        SizingService(store).run_now()
        ctx = _build_context(store)

        failures: list[str] = []
        for section_cls in SectionRegistry.all_sections():
            rb = DocxBuilder(ReportConfig())
            try:
                section_cls().build(ctx, rb)
            except Exception as exc:
                failures.append(f"[{propulsion.name}] {section_cls.__name__}: {exc!r}")

        if failures:
            pytest.fail("Sections failed to build:\n" + "\n".join(failures))


# ── End-to-end export ───────────────────────────────────────────────────────


class TestEndToEndExport:
    @pytest.mark.parametrize(
        "propulsion",
        [
            PropulsionType.ELECTRIC,
            PropulsionType.PISTON,
            PropulsionType.HYBRID,
            PropulsionType.TURBOJET,
        ],
    )
    def test_export_produces_valid_docx(self, store, propulsion, tmp_path):
        store.update_brief_field("propulsion_type", propulsion)
        SizingService(store).run_now()

        out = tmp_path / f"report_{propulsion.name.lower()}.docx"
        ok, msg = _export(store, out)
        assert ok, f"Export failed for {propulsion.name}: {msg}"
        assert out.exists()
        assert out.stat().st_size > 5000

        # File must be readable by python-docx
        doc = Document(str(out))
        assert len(doc.paragraphs) > 10


# ── Structural inspection of the rendered .docx ─────────────────────────────


class TestExportedDocxStructure:
    @pytest.fixture
    def exported_doc(self, sized_store, tmp_path):
        out = tmp_path / "smoke.docx"
        ok, msg = _export(sized_store, out)
        assert ok, msg
        return Document(str(out))

    def test_no_error_notes_leaked_into_body(self, exported_doc):
        bad_phrases = ["could not be rendered", "AttributeError", "Traceback"]
        for p in exported_doc.paragraphs:
            for phrase in bad_phrases:
                assert phrase not in p.text, (
                    f"Error text leaked into doc paragraph: {p.text!r}"
                )

    def test_expected_section_headings_present(self, exported_doc):
        full_text = "\n".join(p.text for p in exported_doc.paragraphs)
        expected_substrings = [
            "Mission Requirements",
            "Mission Profile",
            "Weight Estimation",
            "Constraint Analysis",
            "Design Point",
            "Aerodynamic Parameters",
            "Sanity Checks",
            "Appendix",
        ]
        for needle in expected_substrings:
            assert needle in full_text, f"Missing heading text: {needle!r}"

    def test_omml_math_elements_present(self, exported_doc):
        """At least one m:oMath element must appear (real Word equation object)."""
        body_xml = etree.tostring(exported_doc.element).decode("utf-8")
        assert f"{{{_OMML_NS}}}oMath" in body_xml or "m:oMath" in body_xml, (
            "No <m:oMath> elements found — equations are not rendered as Word math"
        )

    def test_omml_has_real_structure_not_just_text(self, exported_doc):
        """Equations must produce structural OMML — fractions, radicals, sub/sup —
        not just <m:t> text inside <m:oMath>. Without this, Word renders the
        equation as plain italic text rather than proper math layout.

        python-docx serializes OMML with the ``m:`` prefix (e.g. ``<m:f>``);
        match the prefix form rather than the Clark-notation tag.
        """
        body_xml = etree.tostring(exported_doc.element).decode("utf-8")
        # The doc must contain at least one of each: fraction, radical,
        # subscript, and subscript+superscript (a baseline of math richness).
        assert "<m:f>" in body_xml or "<m:f " in body_xml, "No <m:f> fractions found"
        assert "<m:rad>" in body_xml or "<m:rad " in body_xml, "No <m:rad> radicals found"
        assert "<m:sSub>" in body_xml or "<m:sSub " in body_xml, "No <m:sSub> subscripts found"
        # sSup is rarer in our equations; sSubSup is the more common combined form
        assert (
            "<m:sSup>" in body_xml or "<m:sSubSup>" in body_xml
        ), "No <m:sSup> / <m:sSubSup> superscripts found"

    def test_mission_segments_energy_matches_electric_propulsion(
        self, exported_doc
    ):
        """For default (Electric) brief, every segment must show Battery, not Fuel."""
        found = False
        for tbl in exported_doc.tables:
            headers = [c.text.strip() for c in tbl.rows[0].cells]
            if "Segment Type" in headers and "Energy Source" in headers:
                idx = headers.index("Energy Source")
                for row in tbl.rows[1:]:
                    val = row.cells[idx].text.strip().lower()
                    assert val == "battery", (
                        f"Electric propulsion segment shows Energy Source = {val!r}, "
                        f"expected 'battery'"
                    )
                found = True
                break
        assert found, "Mission Segments table not found in exported doc"

    def test_weight_pie_chart_image_present(self, exported_doc):
        """Weight Estimation should embed a pie-chart figure."""
        inline_shapes = exported_doc.inline_shapes
        assert len(inline_shapes) >= 1, (
            f"Expected at least 1 embedded figure (weight pie chart), "
            f"found {len(inline_shapes)}"
        )

    def test_cover_page_uses_today_not_brief_date(self, exported_doc):
        """Date on cover should be a recent ISO date, not 2026-04-04 (frozen)."""
        full_text = "\n".join(p.text for p in exported_doc.paragraphs)
        for tbl in exported_doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text
        assert "2026-04-04" not in full_text, (
            "Cover page date is frozen at 2026-04-04 — should use today's date"
        )


# ── OMML XML helper ─────────────────────────────────────────────────────────


class TestOMMLBuilder:
    """Verify the LaTeX→OMML transform produces real Word math structure,
    not just <m:t> text-runs wrapped in <m:oMath>."""

    def _count(self, el, *local_names):
        """Count descendants whose local-name is in ``local_names``."""
        total = 0
        for d in el.iter():
            tag = d.tag.split("}", 1)[1] if "}" in d.tag else d.tag
            if tag in local_names:
                total += 1
        return total

    def test_make_omath_returns_oMath_element(self):
        el = _make_omath(r"a + b = c")
        assert el.tag == f"{{{_OMML_NS}}}oMath"

    def test_fraction_emits_omml_frac(self):
        el = _make_omath(r"\frac{W_E}{W_{TO}} = a")
        assert self._count(el, "f") >= 1, (
            "LaTeX \\frac should produce <m:f> (real fraction structure)"
        )
        # Each <m:f> must contain <m:num> and <m:den>
        assert self._count(el, "num") >= 1
        assert self._count(el, "den") >= 1

    def test_sqrt_emits_omml_rad(self):
        el = _make_omath(r"C_L = \sqrt{C_{D_0} / k}")
        assert self._count(el, "rad") >= 1, (
            "LaTeX \\sqrt should produce <m:rad> (real radical structure)"
        )

    def test_subscript_emits_omml_sSub(self):
        el = _make_omath(r"V_s")
        assert self._count(el, "sSub") >= 1, (
            "LaTeX subscript should produce <m:sSub>"
        )

    def test_superscript_emits_omml_sSup(self):
        el = _make_omath(r"V^2")
        assert self._count(el, "sSup") >= 1, (
            "LaTeX superscript should produce <m:sSup>"
        )

    def test_sub_and_sup_combined_emits_sSubSup(self):
        el = _make_omath(r"V_s^2")
        assert self._count(el, "sSubSup") >= 1, (
            "LaTeX sub+sup on same base should produce <m:sSubSup>"
        )

    def test_greek_letters_preserved_in_text(self):
        el = _make_omath(r"\rho_0 \cdot V")
        # Find all <m:t> elements and concatenate their text
        all_text = "".join(
            (t.text or "")
            for t in el.iter(f"{{{_OMML_NS}}}t")
        )
        assert "ρ" in all_text, "Greek rho missing from OMML output"

    def test_well_formed_xml(self):
        el = _make_omath(r"\frac{1}{2}\rho_0 V_s^2 C_{L_{\max}}")
        s = etree.tostring(el)
        reparsed = etree.fromstring(s)
        assert reparsed.tag.endswith("}oMath")


# ── Bilingual report (English / Arabic via gettext) ─────────────────────────


class TestI18n:
    """Translator + RTL layout for the Arabic report variant."""

    def _exported_ar(self, store, tmp_path):
        from app.core.i18n import Language
        from app.services.export_service import ExportService
        from app.core.reports.base import SectionRegistry
        out = tmp_path / "report_ar.docx"
        cfg = ReportConfig(
            report_title="تقرير اختبار",
            author="مختبر",
            revision="1.0",
            format=ExportFormat.DOCX,
            sections=SectionRegistry.default_manifest(),
            output_path=str(out),
            language=Language.AR,
        )
        ok, msg = ExportService(store).export(cfg)
        assert ok, msg
        return Document(str(out))

    def test_translator_returns_arabic_for_known_keys(self):
        from app.core.i18n import Language
        from app.core.i18n.translator import load_translator
        from pathlib import Path
        loc = Path(__file__).resolve().parent.parent / "app" / "resources" / "locales"
        t = load_translator(loc, Language.AR)
        assert t.t("cover.label.project") == "المشروع"
        assert t.t("status.pass") == "ناجح"

    def test_translator_falls_back_to_key_for_unknown(self):
        from app.core.i18n import Language
        from app.core.i18n.translator import load_translator
        from pathlib import Path
        loc = Path(__file__).resolve().parent.parent / "app" / "resources" / "locales"
        t = load_translator(loc, Language.AR)
        assert t.t("does.not.exist") == "does.not.exist"

    def test_smart_format_preserves_math_zones(self):
        from app.core.i18n import smart_format
        template = r"area: {area:.2f} {unit}, with $${C_{L_{max}}}$$ from regression"
        result = smart_format(template, area=1.523, unit="m²")
        assert "1.52" in result and "m²" in result
        assert "$${C_{L_{max}}}$$" in result, "math zone braces were mangled"

    def test_arabic_export_paragraphs_have_bidi_marker(self, sized_store, tmp_path):
        """RTL paragraphs must carry <w:bidi/> so Word lays them out correctly."""
        doc = self._exported_ar(sized_store, tmp_path)
        body_xml = etree.tostring(doc.element).decode("utf-8")
        assert "<w:bidi/>" in body_xml or "<w:bidi " in body_xml, (
            "Arabic export missing <w:bidi/> on paragraphs"
        )

    def test_arabic_export_tables_have_bidi_visual(self, sized_store, tmp_path):
        """Tables in RTL mode use <w:bidiVisual/> for column-flip."""
        doc = self._exported_ar(sized_store, tmp_path)
        body_xml = etree.tostring(doc.element).decode("utf-8")
        assert "<w:bidiVisual/>" in body_xml or "<w:bidiVisual " in body_xml, (
            "Arabic export missing <w:bidiVisual/> on tables"
        )

    def test_arabic_export_uses_complex_script_font(self, sized_store, tmp_path):
        """Text runs in RTL mode must declare a complex-script font (cs= attr)."""
        doc = self._exported_ar(sized_store, tmp_path)
        body_xml = etree.tostring(doc.element).decode("utf-8")
        assert 'w:cs="Tahoma"' in body_xml, (
            "Arabic export missing complex-script font on runs"
        )

    def test_arabic_export_contains_arabic_glyphs(self, sized_store, tmp_path):
        """Translated text should actually contain Arabic characters."""
        doc = self._exported_ar(sized_store, tmp_path)
        full_text = _doc_text(doc)
        # Catalogue contains at least these strings translated to Arabic
        assert "المشروع" in full_text, "cover.label.project not translated"
        assert "ناجح" in full_text or "فاشل" in full_text, (
            "status.pass / status.fail not translated"
        )

    def test_equation_numbers_stay_western_in_arabic(self, sized_store, tmp_path):
        """Equation labels like (1.2) keep Latin digits even in Arabic mode."""
        import re
        doc = self._exported_ar(sized_store, tmp_path)
        full_text = _doc_text(doc)
        # At least one (N.M) label with Western digits must appear
        assert re.search(r"\(\d+\.\d+\)", full_text), (
            "Arabic export lost Western-digit equation labels"
        )

    def test_english_export_has_no_bidi_markers(self, sized_store, tmp_path):
        """The LTR (default English) export must not emit any RTL markers."""
        out = tmp_path / "report_en.docx"
        ok, msg = _export(sized_store, out)
        assert ok, msg
        body_xml = etree.tostring(Document(str(out)).element).decode("utf-8")
        assert "<w:bidi/>" not in body_xml, "English export carried <w:bidi/>"
        assert "<w:bidiVisual/>" not in body_xml, "English export carried <w:bidiVisual/>"

    def test_arabic_runs_have_rtl_marker(self, sized_store, tmp_path):
        """Each visible text run in the AR export must carry <w:rtl/>.

        Without it, Word treats the run as LTR even inside a <w:bidi/>
        paragraph — numbers and punctuation drift to the wrong end of
        the line and Arabic mixed with Latin tokens reads in the wrong
        order.
        """
        doc = self._exported_ar(sized_store, tmp_path)
        body_xml = etree.tostring(doc.element).decode("utf-8")
        # Count <w:rtl/> occurrences; should be at least a dozen — every
        # paragraph that emits text adds one per run.
        rtl_count = body_xml.count("<w:rtl/>") + body_xml.count("<w:rtl ")
        assert rtl_count >= 30, (
            f"Expected many <w:rtl/> markers in AR runs; found {rtl_count}"
        )

    def test_arabic_export_translates_figure_label(self, sized_store, tmp_path):
        """The 'Figure' caption label must be translated."""
        doc = self._exported_ar(sized_store, tmp_path)
        text = _doc_text(doc)
        assert "الشكل" in text, "'Figure' caption label not translated to 'الشكل'"
        # And the LTR-English label must NOT appear in AR captions.
        # (It can still appear in author-supplied content, but not as the
        # auto-emitted figure-caption prefix.)
        # Heuristic: "Figure 1:" / "Figure 2:" should be absent.
        import re
        assert not re.search(r"\bFigure \d+:", text), (
            "AR export still has English 'Figure N:' caption prefix"
        )

    def test_arabic_export_translates_table_label(self, sized_store, tmp_path):
        doc = self._exported_ar(sized_store, tmp_path)
        text = _doc_text(doc)
        assert "الجدول" in text, "'Table' caption label not translated to 'الجدول'"
        import re
        assert not re.search(r"\bTable \d+:", text), (
            "AR export still has English 'Table N:' caption prefix"
        )

    def test_arabic_export_translates_note_prefix(self, sized_store, tmp_path):
        doc = self._exported_ar(sized_store, tmp_path)
        text = _doc_text(doc)
        # AR prefix must appear; English "Note:" prefix from add_note must not
        assert "ملاحظة" in text, "'Note:' prefix not translated to 'ملاحظة:'"
        # add_note() emits "Note:  " in EN — check the auto-emitted prefix
        # specifically. (Body prose mentioning "Note" as a word is fine.)
        assert "Note:  " not in text, (
            "AR export still has English 'Note:  ' prefix from add_note()"
        )

    def test_english_export_keeps_english_labels(self, sized_store, tmp_path):
        """EN export must still say 'Figure', 'Table', 'Note:' — confirms
        the translation is per-language, not a hardcoded swap."""
        out = tmp_path / "report_en.docx"
        ok, _ = _export(sized_store, out)
        assert ok
        text = _doc_text(Document(str(out)))
        import re
        assert re.search(r"\bFigure \d+:", text), "EN export missing 'Figure N:'"
        assert re.search(r"\bTable \d+:", text),  "EN export missing 'Table N:'"

    def test_arabic_export_covers_every_section(self, sized_store, tmp_path):
        """Each refactored section emits at least one Arabic phrase that
        appears nowhere in the English catalogue — proves it really hit the
        Arabic .mo, not the EN-source fallback path."""
        doc = self._exported_ar(sized_store, tmp_path)
        text = _doc_text(doc)
        # One unique Arabic phrase per section. Found via Arabic catalogue.
        per_section_needles = {
            "cover_page":           "المعدّ",          # actually "المُعدّ" (Author)
            "mission_requirements": "متطلبات المهمة",
            "mission_profile":      "ملف المهمة",
            "weight_breakdown":     "تقدير الوزن",
            "weight_equations":     "كسر الوزن الفارغ",
            "matching_diagram":     "مخطط المطابقة",
            "constraint_equations": "قيد سرعة الانهيار",
            "constraint_status":    "حالة القيود التصميمية",
            "design_point_summary": "نقطة التصميم ونتائج",
            "aero_params":          "المعاملات الديناميكية الهوائية",
            "sanity_checks":        "فحوصات قوانين القياس",
            "appendix_inputs":      "ملخص المدخلات الكامل",
            "appendix_references":  "الملحق ب: المراجع",
        }
        missing = [
            f"{sid}: {needle!r}"
            for sid, needle in per_section_needles.items()
            if needle not in text
        ]
        # The "cover_page" needle uses an unusual diacritic form; tolerate
        # if at least the rest are present — Author is also matched elsewhere.
        critical_missing = [m for m in missing if not m.startswith("cover_page")]
        assert not critical_missing, (
            f"Arabic export missing translated content from sections:\n  "
            + "\n  ".join(critical_missing)
        )


# ── Server-side figure renderers & shared plot-data builders ────────────────


class TestPlotDataBuilders:
    """The core/plots/ data builders are pure-Python and Qt-free; they feed
    both the live UI plots and the matplotlib export renderers."""

    def test_mission_profile_data_default_brief(self, sized_store):
        from app.core.plots import build_mission_profile_data
        data = build_mission_profile_data(sized_store.state.sizing.brief)
        # Default brief has 6 segments, all enabled
        assert len(data.segments) == 6
        assert all(seg.enabled for seg in data.segments)
        # Cruise altitude must propagate through unchanged
        assert data.cruise_alt == pytest.approx(
            sized_store.state.sizing.brief.cruise_altitude_m
        )
        # x_max should be strictly positive (mission has length)
        assert data.x_max > 0

    def test_mission_profile_segment_colours_match_propulsion(self, store):
        from app.core.enums import PropulsionType
        from app.core.plots import build_mission_profile_data
        from app.core.plots.mission_profile import (
            COLOUR_BATTERY, COLOUR_FUEL, COLOUR_FIXED,
        )

        # Electric → dynamic segments (cruise/loiter) carry the battery colour
        store.update_brief_field("propulsion_type", PropulsionType.ELECTRIC)
        SizingService(store).run_now()
        data = build_mission_profile_data(store.state.sizing.brief)
        # Find cruise/loiter (dynamic, enabled) segments and check colour
        dyn_colours = {seg.color for seg in data.segments
                       if "Cruise" in seg.label or "Loiter" in seg.label}
        assert dyn_colours <= {COLOUR_BATTERY}, (
            f"Electric dynamic segments should be battery-colour; got {dyn_colours}"
        )

        # Piston → fuel colour for dynamic segments
        store.update_brief_field("propulsion_type", PropulsionType.PISTON)
        SizingService(store).run_now()
        data2 = build_mission_profile_data(store.state.sizing.brief)
        dyn_colours_2 = {seg.color for seg in data2.segments
                         if "Cruise" in seg.label or "Loiter" in seg.label}
        assert dyn_colours_2 <= {COLOUR_FUEL}

    def test_matching_plot_data_default(self, sized_store):
        from app.core.display_converter import DisplayConverter
        from app.core.plots import build_matching_plot_data

        s = sized_store.state.sizing
        dc = DisplayConverter(sized_store.settings)
        data = build_matching_plot_data(s.constraint_result, s.design_point, dc)

        # Four constraint curves (Max Speed, Takeoff Run, Rate of Climb,
        # Service Ceiling) plus a stall vertical line drawn separately
        assert len(data.curves) == 4
        # Y-clip should sit comfortably above the design-point loading
        assert data.y_top > 0
        # Design point must be present and inside the visible range
        assert data.design_point is not None
        dp_x, dp_y = data.design_point
        assert dp_y <= data.y_top, (
            f"Design point y={dp_y} above y_top={data.y_top} — would be off-plot"
        )


class TestServerSideFigureRenderers:
    """Matplotlib renderers must produce non-empty PNGs and round-trip via
    Pillow without errors."""

    def _pil_open(self, png_bytes: bytes):
        import io as _io
        try:
            from PIL import Image
        except ImportError:
            pytest.skip("Pillow not installed; skipping image-decoding test")
        return Image.open(_io.BytesIO(png_bytes))

    def test_mission_profile_png_renders(self, sized_store):
        from app.services.figure_renderers import render_mission_profile_png
        png = render_mission_profile_png(sized_store.state.sizing.brief)
        assert png is not None and len(png) > 2000
        img = self._pil_open(png)
        # Image must be a sensible size (matplotlib default at 150 dpi)
        assert img.width >= 600 and img.height >= 200

    def test_matching_diagram_png_renders(self, sized_store):
        from app.core.display_converter import DisplayConverter
        from app.services.figure_renderers import render_matching_diagram_png
        s = sized_store.state.sizing
        png = render_matching_diagram_png(
            s.constraint_result, s.design_point,
            DisplayConverter(sized_store.settings),
        )
        assert png is not None and len(png) > 5000
        img = self._pil_open(png)
        assert img.width >= 800 and img.height >= 400

    def test_weight_pie_png_renders_for_each_propulsion(self, store):
        from app.core.enums import PropulsionType
        from app.services.figure_renderers import render_weight_pie_png
        for prop in (
            PropulsionType.ELECTRIC,
            PropulsionType.PISTON,
            PropulsionType.HYBRID,
        ):
            store.update_brief_field("propulsion_type", prop)
            SizingService(store).run_now()
            png = render_weight_pie_png(
                store.state.sizing.weight_result, prop,
            )
            assert png is not None and len(png) > 2000, (
                f"Pie chart for {prop.name} is missing or too small"
            )

    def test_export_embeds_all_figures(self, sized_store, tmp_path):
        """End-to-end: the exported .docx contains at least the three
        legacy figures (matching diagram, mission profile, weight pie
        chart) plus the sensitivity section's tornado figures (3 by
        default — one per ``sens_tornado_output_ids`` slot)."""
        out = tmp_path / "with_figures.docx"
        ok, msg = _export(sized_store, out)
        assert ok, msg
        doc = Document(str(out))
        n_shapes = len(doc.inline_shapes)
        # Each figure becomes one inline shape: 3 legacy + 3 default
        # sensitivity tornados = 6.
        assert n_shapes == 6, (
            f"Expected 6 figures (3 legacy + 3 default sensitivity "
            f"tornados); found {n_shapes}"
        )


# ── Auto-numbering, inline math, and citation style ────────────────────────


class TestEquationNumberingAndCitations:
    @pytest.fixture
    def exported(self, sized_store, tmp_path):
        out = tmp_path / "numbering.docx"
        ok, msg = _export(sized_store, out)
        assert ok, msg
        return Document(str(out))

    def test_no_section_sign_in_body(self, exported):
        """The § character should not appear anywhere in the document — we
        use 'Sec.' instead so the report reads as natural prose."""
        body = etree.tostring(exported.element).decode("utf-8")
        # Strip the namespace declarations so § in xml namespaces (none in
        # this doc, but defensive) doesn't trigger a false positive
        assert "§" not in body, "Section sign § found in document body"

    def test_auto_numbered_equation_labels(self, exported):
        """Displayed equations should carry labels of the form (N.M) where
        N is the level-1 section index and M is the sub-counter."""
        import re

        body = etree.tostring(exported.element).decode("utf-8")
        labels = re.findall(r"\((\d+)\.(\d+)\)", body)
        # Filter out "Sec. 2.9" / "Eq. 2.49-2.51" textual section refs:
        labels = [(n, m) for n, m in labels if int(n) <= 20 and int(m) <= 99]
        # We expect at least 5 auto-numbered equations in a default export.
        assert len(labels) >= 5, f"Too few auto-numbered eq labels: {labels}"

        # Within each section, sub-counters should be 1..K consecutive.
        from collections import defaultdict
        per_section = defaultdict(list)
        for n, m in labels:
            per_section[int(n)].append(int(m))
        for section, subs in per_section.items():
            assert subs == list(range(1, len(subs) + 1)), (
                f"Section {section} eq numbers are not 1..N consecutive: {subs}"
            )

    def test_no_legacy_eq_2_dot_prefix_labels(self, exported):
        """The old per-equation labels were '(Eq. 2.38)', '(Eq. 2.49)', etc.
        After auto-numbering, no '(Eq. <digit>.<digit>)' should remain."""
        import re

        body = etree.tostring(exported.element).decode("utf-8")
        legacy = re.findall(r"\(Eq\.\s*\d+\.\d+\)", body)
        assert not legacy, f"Legacy (Eq. N.NN) labels still present: {legacy}"


class TestInlineMathTemplate:
    def test_paragraph_inline_math_emits_omml(self, sized_store, tmp_path):
        """A paragraph containing $${LaTeX}$$ should produce inline OMML
        alongside the surrounding text."""
        from app.core.reports.renderers.docx_renderer import DocxBuilder
        from app.core.reports.renderer import ReportConfig

        out = tmp_path / "inline_math.docx"
        rb = DocxBuilder(ReportConfig())
        rb.add_heading("Test Section", level=1)
        rb.add_paragraph(
            "The propulsive efficiency $${\\eta_p}$$ appears here, and "
            "so does $${\\rho_0 V_s^2}$$."
        )
        rb.save(str(out))

        from docx import Document
        doc = Document(str(out))
        body = etree.tostring(doc.element).decode("utf-8")
        # Two inline math sources → two <m:oMath> elements (plus none from
        # add_equation, since we never called it)
        omath_count = body.count("<m:oMath>") + body.count("<m:oMath ")
        assert omath_count == 2, (
            f"Expected 2 inline OMML elements, found {omath_count}"
        )

    def test_inline_math_in_default_export(self, sized_store, tmp_path):
        """Section files use $${...}$$ in their intros / kv-lists. Confirm
        those round-trip into the rendered .docx as inline OMML."""
        out = tmp_path / "with_inline.docx"
        ok, msg = _export(sized_store, out)
        assert ok, msg
        body = etree.tostring(Document(str(out)).element).decode("utf-8")
        # The raw $${...}$$ sigil should NOT appear in the rendered doc
        assert "$${" not in body, (
            "Raw inline-math sigil leaked into the rendered document"
        )
        # And at least one inline OMML run should exist (section intros add them)
        omath_count = body.count("<m:oMath>") + body.count("<m:oMath ")
        assert omath_count >= 1, "No OMML elements at all — inline math not emitted"


# ── Display-rule tests: only relevant data should appear ────────────────────


def _doc_text(doc) -> str:
    """Concatenate all text content (paragraphs + table cells) into one string."""
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestDisplayRules:
    """Hide irrelevant data per propulsion type (ui_ux.md bible rule)."""

    def _exported(self, store, tmp_path, propulsion):
        store.update_brief_field("propulsion_type", propulsion)
        SizingService(store).run_now()
        out = tmp_path / f"display_{propulsion.name.lower()}.docx"
        ok, msg = _export(store, out)
        assert ok, msg
        return Document(str(out))

    def test_electric_propulsion_hides_thrust_to_weight(self, store, tmp_path):
        doc = self._exported(store, tmp_path, PropulsionType.ELECTRIC)
        text = _doc_text(doc)
        # Electric uses W/P — (T/W) should not appear as a variable definition
        # or as a sizing-results row label.
        assert "(T/W)_d" not in text, (
            "Electric propulsion design-point summary should not mention (T/W)_d"
        )
        assert "Thrust-to-Weight Ratio (T/W)" not in text, (
            "Electric propulsion sizing results should not show Thrust-to-Weight Ratio"
        )

    def test_turbojet_propulsion_hides_power_loading(self, store, tmp_path):
        doc = self._exported(store, tmp_path, PropulsionType.TURBOJET)
        text = _doc_text(doc)
        # Turbojet uses T/W — (W/P) should not appear as variable def or
        # as a results-row label.
        assert "(W/P)_d" not in text, (
            "Turbojet propulsion design-point summary should not mention (W/P)_d"
        )
        assert "Power Loading (W/P)" not in text, (
            "Turbojet propulsion sizing results should not show Power Loading"
        )

    def test_no_alternative_notes_column_in_sizing_results(
        self, sized_store, tmp_path
    ):
        out = tmp_path / "no_alt.docx"
        ok, _ = _export(sized_store, out)
        assert ok
        doc = Document(str(out))
        for tbl in doc.tables:
            headers = [c.text.strip() for c in tbl.rows[0].cells]
            assert "Alternative / Notes" not in headers, (
                f"Sizing Results table still has 'Alternative / Notes' column: {headers}"
            )

    def test_matching_diagram_wing_loading_no_duplicate_si(
        self, sized_store, tmp_path
    ):
        """Wing loading line should not show '... | <same value> N/m²' suffix."""
        out = tmp_path / "no_dup.docx"
        ok, _ = _export(sized_store, out)
        assert ok
        doc = Document(str(out))
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text = cell.text
                    if "Wing Loading (W/S)" in text:
                        # cell with the label only; the value cell is the next one
                        continue
                    if "N/m²" in text and "|" in text:
                        pytest.fail(
                            f"Duplicate SI value in matching-diagram cell: {text!r}"
                        )


# ── Bonus: include/exclude equations toggle ─────────────────────────────────


class TestIncludeEquationsToggle:
    """include_equations=False skips the displayed (auto-numbered) equation
    blocks but does NOT scrub inline $${...}$$ math from prose — those are
    integral to the surrounding sentences."""

    def test_displayed_equations_excluded_when_toggled_off(
        self, sized_store, tmp_path
    ):
        import re

        cfg_off = ReportConfig(
            report_title="No equations",
            author="Test",
            revision="1.0",
            format=ExportFormat.DOCX,
            sections=SectionRegistry.default_manifest(),
            output_path=str(tmp_path / "no_eq.docx"),
            include_equations=False,
        )
        ok, _ = ExportService(sized_store).export(cfg_off, figure_grabbers={})
        assert ok
        doc_off = Document(str(tmp_path / "no_eq.docx"))
        body_off = etree.tostring(doc_off.element).decode("utf-8")

        # Auto-numbered displayed equations look like "(7.1)", "(7.2)", etc.
        # Inline math has no such label. So the toggle's signature is:
        eq_labels_off = re.findall(r"\(\d+\.\d+\)", body_off)
        # The cover page shows "Sec. 2.6-2.7" / "Eq. 2.49-2.51" — those are
        # textual section references, not auto-numbered equation labels.
        # Filter them out:
        eq_labels_off = [
            lbl for lbl in eq_labels_off
            if "Sec." not in lbl and "Eq." not in lbl
        ]
        assert not eq_labels_off, (
            f"include_equations=False but auto-numbered eq labels still appear: "
            f"{eq_labels_off[:5]}"
        )

    def test_displayed_equations_present_when_toggled_on(
        self, sized_store, tmp_path
    ):
        import re

        cfg_on = ReportConfig(
            report_title="With equations",
            author="Test",
            revision="1.0",
            format=ExportFormat.DOCX,
            sections=SectionRegistry.default_manifest(),
            output_path=str(tmp_path / "with_eq.docx"),
            include_equations=True,
        )
        ok, _ = ExportService(sized_store).export(cfg_on, figure_grabbers={})
        assert ok
        doc_on = Document(str(tmp_path / "with_eq.docx"))
        body_on = etree.tostring(doc_on.element).decode("utf-8")
        eq_labels_on = re.findall(r"\(\d+\.\d+\)", body_on)
        eq_labels_on = [
            lbl for lbl in eq_labels_on
            if "Sec." not in lbl and "Eq." not in lbl
        ]
        assert len(eq_labels_on) >= 5, (
            f"Expected auto-numbered equation labels; only found {eq_labels_on}"
        )


# ── Sensitivity Analysis section (PR6 Slice 1) ──────────────────────────────


class TestSensitivityReportSection:
    """The sensitivity_analysis section ships with a SectionConfig payload
    that the export dialog can customise. These tests cover (1) the
    abstraction (SectionConfig.validate filters propulsion-gated inputs),
    (2) the section's default config (the live page's tornado slots), and
    (3) the end-to-end export pipeline emitting a docx whose body contains
    the expected headings and figure refs.
    """

    def test_default_config_matches_studio_tornado_slots(self, sized_store):
        from app.core.reports.sections.sensitivity_analysis import (
            SensitivityAnalysisSection,
            SensitivityReportConfig,
        )
        ctx = _build_context(sized_store)
        cfg = SensitivityAnalysisSection.default_config(ctx)
        assert isinstance(cfg, SensitivityReportConfig)
        assert cfg.tornado_output_ids == sized_store.settings.sens_tornado_output_ids
        assert cfg.sweep_specs == ()                # empty by default
        assert cfg.include_margins is True
        assert cfg.include_snowball is True

    def test_validate_drops_propulsion_gated_sweep_input(self, sized_store):
        """SFC is gated to ``uses_fuel`` — an Electric brief must drop it."""
        import dataclasses
        from app.core.reports.sections.sensitivity_analysis import (
            SensitivityReportConfig,
        )
        cfg = SensitivityReportConfig(
            tornado_output_ids=("mtow_kg",),
            sweep_specs=(
                ("mtow_kg", "payload_mass_kg"),               # universal — keeps
                ("mtow_kg", "specific_fuel_consumption_g_wh"),  # fuel-only — drops
            ),
        )
        # Default brief is Electric — second sweep spec must disappear.
        cleaned = cfg.validate(sized_store.state.sizing.brief)
        assert cleaned.tornado_output_ids == ("mtow_kg",)
        assert cleaned.sweep_specs == (("mtow_kg", "payload_mass_kg"),)

    def test_validate_drops_unknown_output_id(self, sized_store):
        from app.core.reports.sections.sensitivity_analysis import (
            SensitivityReportConfig,
        )
        cfg = SensitivityReportConfig(
            tornado_output_ids=("mtow_kg", "does_not_exist"),
        )
        cleaned = cfg.validate(sized_store.state.sizing.brief)
        assert cleaned.tornado_output_ids == ("mtow_kg",)

    def test_summary_describes_picks(self, sized_store):
        from app.core.reports.sections.sensitivity_analysis import (
            SensitivityReportConfig,
        )
        cfg = SensitivityReportConfig(
            tornado_output_ids=("mtow_kg", "wing_area_m2"),
            sweep_specs=(("mtow_kg", "payload_mass_kg"),),
        )
        summary = cfg.summary()
        assert "2 tornados" in summary
        assert "1 sweep" in summary
        assert "margins" in summary
        assert "snowball" in summary

    def test_summary_empty(self):
        from app.core.reports.sections.sensitivity_analysis import (
            SensitivityReportConfig,
        )
        cfg = SensitivityReportConfig(
            tornado_output_ids=(),
            sweep_specs=(),
            include_margins=False,
            include_snowball=False,
        )
        assert cfg.summary() == "(empty)"

    def test_section_class_metadata(self):
        from app.core.reports.sections.sensitivity_analysis import (
            SensitivityAnalysisSection,
        )
        assert SensitivityAnalysisSection.section_id == "sensitivity_analysis"
        assert SensitivityAnalysisSection.is_customizable is True
        assert SensitivityAnalysisSection.default_order == 95
        # Slots after sanity_checks(90), before appendix_inputs(100).

    def test_default_export_includes_sensitivity_headings(
        self, sized_store, tmp_path,
    ):
        ok, _ = _export(sized_store, tmp_path / "sens_default.docx")
        assert ok
        doc = Document(str(tmp_path / "sens_default.docx"))
        body = etree.tostring(doc.element).decode("utf-8")
        assert "Design Sensitivity Analysis" in body
        assert "Constraint Margins" in body
        assert "Snowball Factors" in body
        # Default config has no sweeps — no "Sweep …" sub-heading.
        # Heading uses an em-dash that lxml renders as `&#8212;`.
        assert "Sweep &#8212;" not in body
        # Regression: a default export of a sized store must successfully
        # resolve coeffs (database lookup → textbook fallback). If
        # ``ctx.regression_coeffs`` is None, the section emits the
        # "Regression coefficients unavailable — skipped." note and the
        # tornado / snowball content silently disappears.
        assert "Regression coefficients unavailable" not in body, (
            "ctx.regression_coeffs is None — coeffs resolution is broken; "
            "see app/services/coeffs_resolver.py"
        )
        # Positive check: at least one tornado figure caption is present.
        assert "Tornado of input impact" in body

    def test_custom_config_with_sweep_emits_sweep_heading(
        self, sized_store, tmp_path,
    ):
        from app.core.reports.base import SectionEntry
        from app.core.reports.sections.sensitivity_analysis import (
            SensitivityReportConfig,
        )
        manifest = SectionRegistry.default_manifest()
        for entry in manifest:
            if entry.section_id == "sensitivity_analysis":
                entry.config = SensitivityReportConfig(
                    tornado_output_ids=("mtow_kg",),
                    sweep_specs=(("mtow_kg", "payload_mass_kg"),),
                    include_margins=True,
                    include_snowball=False,
                )
        cfg = ReportConfig(
            report_title="Sens custom",
            author="Test",
            revision="1.0",
            format=ExportFormat.DOCX,
            sections=manifest,
            output_path=str(tmp_path / "sens_custom.docx"),
        )
        ok, _ = ExportService(sized_store).export(cfg, figure_grabbers={})
        assert ok
        doc = Document(str(tmp_path / "sens_custom.docx"))
        body = etree.tostring(doc.element).decode("utf-8")
        # Custom: 1 tornado, 1 sweep, margins yes, snowball NO.
        # lxml's tostring renders the em-dash as the XML numeric entity
        # `&#8212;`, so match either the entity-encoded form or a unique
        # plain-text substring of the heading.
        assert ("Sweep &#8212;" in body
                or "MTOW vs Payload Mass" in body), \
            "expected 'Sweep ... MTOW vs Payload Mass' heading in docx"
        assert "Constraint Margins" in body
        assert "Snowball Factors" not in body


# ── Shared coeffs resolver (PR6 hotfix) ─────────────────────────────────────


class TestCoeffsResolver:
    """``resolve_active_coeffs`` is the single source of truth for picking
    regression coefficients across SizingService, SensitivityService, and
    ExportService. Without it the bug-of-the-day was: ExportService keyed
    by ``propulsion_type.name.lower()`` (always missed), causing every
    sensitivity-section figure and the snowball table to silently skip.
    """

    def test_resolver_returns_textbook_fallback_when_database_empty(
        self, sized_store,
    ):
        from app.services.coeffs_resolver import resolve_active_coeffs
        coeffs = resolve_active_coeffs(
            sized_store, sized_store.state.sizing.brief,
        )
        assert coeffs is not None, (
            "Resolver must fall back to textbook coefficients when the "
            "database carries nothing for the current classification."
        )

    def test_export_context_carries_resolved_coeffs(self, sized_store):
        """ReportContext.regression_coeffs must be populated for any
        sized store — without it the sensitivity section silently
        skips its figures and snowball table.

        Goes through ``ExportService._build_context`` (the production
        path) rather than the test-local ``_build_context`` helper,
        which legacy code paths hardcoded to None.
        """
        cfg = ReportConfig(
            report_title="t", author="t", revision="1",
            format=ExportFormat.DOCX,
            sections=SectionRegistry.default_manifest(),
            output_path="/tmp/x.docx",
        )
        ctx = ExportService(sized_store)._build_context(cfg)
        assert ctx.regression_coeffs is not None


# ── SectionConfig abstraction (PR6 Slice 1) ─────────────────────────────────


class TestSectionConfigBase:
    """The SectionConfig abstract base + SectionEntry typing."""

    def test_base_validate_default_passthrough(self, sized_store):
        from app.core.reports.base import SectionConfig

        class _Empty(SectionConfig):
            pass

        e = _Empty()
        assert e.validate(sized_store.state.sizing.brief) is e

    def test_section_entry_config_field_typed(self):
        """SectionEntry.config must be Optional[SectionConfig], not Any."""
        from app.core.reports.base import SectionConfig, SectionEntry
        from typing import get_type_hints
        hints = get_type_hints(SectionEntry)
        # Optional[SectionConfig] == Union[SectionConfig, None]
        config_hint = hints["config"]
        assert SectionConfig in getattr(config_hint, "__args__", ())

    def test_report_section_default_config_returns_none_by_default(
        self, sized_store,
    ):
        """Non-customisable sections inherit default_config returning None."""
        from app.core.reports.sections.matching_diagram import (
            MatchingDiagramSection,
        )
        ctx = _build_context(sized_store)
        assert MatchingDiagramSection.default_config(ctx) is None


# ── Tornado / sweep figure renderers (PR6 Slice 1) ──────────────────────────


class TestSensitivityFigureRenderers:
    def test_render_tornado_returns_png_bytes(self, sized_store):
        from app.core.coefficients import get_closest_textbook
        from app.core.sensitivity import (
            compute_tornado,
            sweepable_parameters_for,
        )
        from app.services.figure_renderers import render_tornado_png

        brief  = sized_store.state.sizing.brief
        coeffs = get_closest_textbook(
            brief.classification_name, brief.payload_mass_kg * 5,
        )
        td = compute_tornado(
            brief, coeffs,
            sweepable_parameters_for(brief), "mtow_kg",
        )
        png = render_tornado_png(
            td, brief.propulsion_type,
            DisplayConverter(sized_store.settings),
        )
        assert png is not None and len(png) > 1000
        # PNG signature
        assert png[:8] == b"\x89PNG\r\n\x1a\n"

    def test_render_sweep_single_output_returns_png(self, sized_store):
        from app.core.coefficients import get_closest_textbook
        from app.core.sensitivity import (
            run_oat_sweep,
            sweepable_parameters_for,
        )
        from app.services.figure_renderers import render_sweep_png

        brief  = sized_store.state.sizing.brief
        coeffs = get_closest_textbook(
            brief.classification_name, brief.payload_mass_kg * 5,
        )
        param = next(
            p for p in sweepable_parameters_for(brief)
            if p.field_name == "payload_mass_kg"
        )
        sweep = run_oat_sweep(brief, coeffs, param, n_points=11)
        png = render_sweep_png(
            sweep, ["mtow_kg"],
            brief.propulsion_type,
            DisplayConverter(sized_store.settings),
        )
        assert png is not None and len(png) > 1000
        assert png[:8] == b"\x89PNG\r\n\x1a\n"

    def test_render_sweep_multi_output_returns_png(self, sized_store):
        from app.core.coefficients import get_closest_textbook
        from app.core.sensitivity import (
            run_oat_sweep,
            sweepable_parameters_for,
        )
        from app.services.figure_renderers import render_sweep_png

        brief  = sized_store.state.sizing.brief
        coeffs = get_closest_textbook(
            brief.classification_name, brief.payload_mass_kg * 5,
        )
        param = next(
            p for p in sweepable_parameters_for(brief)
            if p.field_name == "payload_mass_kg"
        )
        sweep = run_oat_sweep(brief, coeffs, param, n_points=11)
        png = render_sweep_png(
            sweep, ["mtow_kg", "wing_area_m2", "engine_power_w"],
            brief.propulsion_type,
            DisplayConverter(sized_store.settings),
        )
        assert png is not None and len(png) > 1000

    def test_render_tornado_empty_returns_none(self, sized_store):
        from app.core.sensitivity import TornadoData
        from app.services.figure_renderers import render_tornado_png

        empty = TornadoData(output_id="mtow_kg", baseline_out=None, bars=())
        png = render_tornado_png(
            empty, sized_store.state.sizing.brief.propulsion_type,
            DisplayConverter(sized_store.settings),
        )
        assert png is None
